"""
Script para gerar documento .docx com a formulação matemática
detalhada das Análises 3 e 5.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

def add_equation(doc, text, bold=False):
    """Adiciona um parágrafo de equação centralizado."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p

def add_body(doc, text):
    """Adiciona parágrafo de corpo."""
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    return p

def add_bullet(doc, text):
    """Adiciona item de lista."""
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def set_margins(doc, top=2.5, bottom=2.5, left=3.0, right=2.5):
    """Define margens do documento em cm."""
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)

def gerar_documento():
    doc = Document()
    set_margins(doc)
    
    # Estilo Normal
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    
    # ========================================
    # TÍTULO
    # ========================================
    title = doc.add_heading('Formulação Matemática das Análises de Otimização de Microrredes', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    add_body(doc, 'Este documento apresenta a formulação matemática completa das Análises 3 (Heurística com Deslizamento de Carga) e 5 (MILP com Deslizamento Integrado), detalhando todas as variáveis, restrições, funções objetivo e algoritmos utilizados no sistema de gerenciamento de microrredes.')
    
    # ========================================
    # 1. NOMENCLATURA
    # ========================================
    doc.add_heading('1. Nomenclatura e Definições', level=1)
    
    doc.add_heading('1.1 Conjuntos e Índices', level=2)
    
    table = doc.add_table(rows=1, cols=2, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = 'Símbolo'
    hdr[1].text = 'Descrição'
    for cell in hdr:
        for paragraph in cell.paragraphs:
            paragraph.runs[0].bold = True
    
    dados_conjuntos = [
        ('t ∈ T = {0, 1, ..., 1439}', 'Períodos de tempo (minutos do dia), |T| = 1440'),
        ('k ∈ K', 'Conjunto de cargas flexíveis (prioridade 2 e 4)'),
        ('s ∈ Sₖ', 'Conjunto de possíveis horários de início da carga k (passo de 15 min)'),
        ('F = {sol, bat, die, bio, conc}', 'Conjunto de fontes de energia disponíveis'),
    ]
    for simb, desc in dados_conjuntos:
        row = table.add_row().cells
        row[0].text = simb
        row[1].text = desc
    
    doc.add_heading('1.2 Parâmetros (Dados de Entrada)', level=2)
    
    table2 = doc.add_table(rows=1, cols=3, style='Table Grid')
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2 = table2.rows[0].cells
    hdr2[0].text = 'Símbolo'
    hdr2[1].text = 'Unidade'
    hdr2[2].text = 'Descrição'
    for cell in hdr2:
        for paragraph in cell.paragraphs:
            paragraph.runs[0].bold = True
    
    params = [
        ('D(t)', 'kW', 'Demanda de carga no período t'),
        ('D_fixo(t)', 'kW', 'Demanda fixa (cargas com prioridade 1 e 3)'),
        ('G_sol(t)', 'kW', 'Geração solar disponível no período t'),
        ('P_bat_max', 'kW', 'Potência máxima de carga/descarga da bateria'),
        ('P_die_max', 'kW', 'Potência máxima do gerador diesel'),
        ('E_bat_max', 'kWh', 'Capacidade máxima da bateria (SOC_max)'),
        ('E_bat_min', 'kWh', 'Capacidade mínima da bateria (SOC_min)'),
        ('E_bat_0', 'kWh', 'Estado inicial de carga da bateria'),
        ('η', '-', 'Eficiência da bateria (0 a 1)'),
        ('E_die_0', 'L', 'Nível inicial do tanque diesel'),
        ('E_die_max', 'L', 'Capacidade máxima do tanque diesel'),
        ('σ', 'L/kWh', 'Consumo específico do diesel (≈ 0.2)'),
        ('c_sol', 'R$/kWh', 'Custo operacional solar (manutenção)'),
        ('c_bat', 'R$/kWh', 'Custo de uso da bateria'),
        ('c_die', 'R$/kWh', 'Custo do diesel por kWh'),
        ('c_conc', 'R$/kWh', 'Tarifa da concessionária'),
        ('c_startup', 'R$', 'Custo de inicialização do diesel (= 50)'),
        ('c_curtail', 'R$/kWh', 'Penalidade por desperdiçar solar (= 10)'),
        ('c_suav', 'R$/kWh', 'Penalidade de suavização da bateria (= 0.5)'),
        ('pₖ', 'kW', 'Potência da carga flexível k'),
        ('dₖ', 'min', 'Duração da carga flexível k'),
    ]
    for simb, unid, desc in params:
        row = table2.add_row().cells
        row[0].text = simb
        row[1].text = unid
        row[2].text = desc
    
    # ========================================
    # 2. ANÁLISE 3 — HEURÍSTICA
    # ========================================
    doc.add_page_break()
    doc.add_heading('2. Análise 3 — Otimização Heurística com Deslizamento de Carga', level=1)
    
    doc.add_heading('2.1 Visão Geral', level=2)
    add_body(doc, 'A Análise 3 utiliza um algoritmo heurístico de despacho em cascata combinado com deslizamento de cargas. O processo ocorre em três etapas sequenciais:')
    add_bullet(doc, 'Etapa 1: Simulação com curva de carga original para obter perfil de custo')
    add_bullet(doc, 'Etapa 2: Deslizamento de cargas flexíveis (prioridade 2 e 4) para períodos de menor custo')
    add_bullet(doc, 'Etapa 3: Nova simulação com curva de carga otimizada')
    
    add_body(doc, 'A principal característica desta análise é o despacho em cascata: as fontes são ordenadas por custo crescente (R$/kWh) e utilizadas sequencialmente até suprir toda a demanda. A concessionária atua como fonte de backup para qualquer déficit residual.')
    
    doc.add_heading('2.2 Algoritmo de Despacho em Cascata', level=2)
    
    add_body(doc, 'Para cada período t = 0, 1, ..., 1439:')
    
    doc.add_heading('2.2.1 Ordenação de Fontes por Custo', level=3)
    add_body(doc, 'As fontes disponíveis são ordenadas pelo custo por kWh em ordem crescente:')
    add_equation(doc, 'Ordem = sort_ascending({c_sol, c_bat, c_bio, c_die})')
    add_body(doc, 'Tipicamente: Solar → Bateria → Biogás → Diesel → Concessionária')
    
    doc.add_heading('2.2.2 Despacho Solar', level=3)
    add_body(doc, 'A geração solar é a primeira fonte utilizada. O excesso é encaminhado para recarga da bateria:')
    add_equation(doc, 'P_sol(t) = min(G_sol(t), D_restante(t))')
    add_equation(doc, 'D_restante(t) ← D_restante(t) − P_sol(t)')
    add_equation(doc, 'Excesso_sol(t) = G_sol(t) − P_sol(t)')
    add_equation(doc, 'Custo_sol(t) = P_sol(t) × c_sol / 60')
    
    doc.add_heading('2.2.3 Gerenciamento da Bateria', level=3)
    add_body(doc, 'A bateria possui dois modos de operação mutuamente exclusivos: descarga (suprir demanda) e recarga (armazenar excesso solar).')
    
    add_body(doc, 'Descarga (se D_restante(t) > 0 e E_bat(t) > E_bat_min):')
    add_equation(doc, 'E_disp = E_bat(t) − E_bat_min     [kWh disponível acima do mínimo]')
    add_equation(doc, 'P_max_energia = E_disp × 60 × η     [kW máximo entregável em 1 min]')
    add_equation(doc, 'P_desc(t) = min(P_bat_max, D_restante(t), P_max_energia)')
    add_equation(doc, 'E_retirada = P_desc(t) / (60 × η)     [kWh retirados do armazenamento]')
    add_equation(doc, 'E_bat(t+1) = E_bat(t) − E_retirada')
    add_equation(doc, 'D_restante(t) ← D_restante(t) − P_desc(t)')
    
    add_body(doc, 'Recarga (se Excesso_sol(t) > 0 e E_bat(t) < E_bat_max):')
    add_equation(doc, 'Espaço = E_bat_max − E_bat(t)     [kWh até o máximo]')
    add_equation(doc, 'P_max_espaço = Espaço × 60 / η     [kW máximo absorvível em 1 min]')
    add_equation(doc, 'P_rec(t) = min(P_bat_max, Excesso_sol(t), P_max_espaço)')
    add_equation(doc, 'E_armazenada = P_rec(t) × η / 60     [kWh efetivamente armazenados]')
    add_equation(doc, 'E_bat(t+1) = E_bat(t) + E_armazenada')
    
    add_body(doc, 'Limites operacionais da bateria:')
    add_equation(doc, 'E_bat_min = Capacidade × Capacidade_min / 100')
    add_equation(doc, 'E_bat_max = Capacidade × Capacidade_max / 100')
    add_equation(doc, 'E_bat_min ≤ E_bat(t) ≤ E_bat_max,   ∀t ∈ T')
    
    doc.add_heading('2.2.4 Despacho do Biogás', level=3)
    add_body(doc, 'O biogás possui regeneração contínua e é despachado conforme disponibilidade:')
    add_equation(doc, 'Se E_bio(t) < E_bio_max:  E_bio(t) ← E_bio(t) + G_bio')
    add_equation(doc, 'P_bio(t) = min(P_bio_max, D_restante(t))')
    add_equation(doc, 'D_restante(t) ← D_restante(t) − P_bio(t)')
    add_equation(doc, 'E_bio(t+1) = E_bio(t) − Consumo_bio(P_bio(t))')
    
    doc.add_heading('2.2.5 Despacho do Diesel', level=3)
    add_body(doc, 'O gerador diesel opera sob demanda com consumo do tanque:')
    add_equation(doc, 'P_die(t) = min(P_die_max, D_restante(t))')
    add_equation(doc, 'D_restante(t) ← D_restante(t) − P_die(t)')
    add_equation(doc, 'E_die(t+1) = E_die(t) − Consumo_die(P_die(t))')
    add_equation(doc, 'Custo_die(t) = P_die(t) × c_die / 60')
    
    doc.add_heading('2.2.6 Concessionária (Backup)', level=3)
    add_body(doc, 'A concessionária supre qualquer déficit energético restante:')
    add_equation(doc, 'P_conc(t) = D_restante(t)     [todo déficit residual]')
    add_equation(doc, 'Custo_conc(t) = P_conc(t) × c_conc / 60')
    
    doc.add_heading('2.2.7 Custo Total Instantâneo', level=3)
    add_equation(doc, 'C(t) = Custo_sol(t) + Custo_bat(t) + Custo_bio(t) + Custo_die(t) + Custo_conc(t)')
    
    doc.add_heading('2.3 Deslizamento de Cargas', level=2)
    add_body(doc, 'Após a Etapa 1, a curva de custo instantâneo C(t) é utilizada para deslizar cargas flexíveis (prioridade 2 e 4) para períodos de menor custo. O processo é:')
    
    add_bullet(doc, '1. Identifica cargas com prioridade 2 e 4 (flexíveis)')
    add_bullet(doc, '2. Calcula o custo médio por hora a partir de C(t)')
    add_bullet(doc, '3. Para cada carga flexível k com duração dₖ:')
    add_bullet(doc, '   a) Busca a janela contígua de dₖ minutos com menor custo acumulado')
    add_bullet(doc, '   b) Move a carga para essa janela')
    add_bullet(doc, '4. Reconstrói a curva de carga com as cargas deslocadas')
    
    add_body(doc, 'A curva de carga otimizada é então:')
    add_equation(doc, 'D_otim(t) = D_fixo(t) + Σₖ pₖ × 𝟙[t ∈ janela_otimizada(k)]')
    
    add_body(doc, 'Onde 𝟙[·] é a função indicadora que vale 1 se t pertence à nova janela de operação da carga k.')
    
    doc.add_heading('2.4 Custo Total da Análise 3', level=2)
    add_equation(doc, 'Custo_Total_A3 = Σ_{t=0}^{1439} C(t)')
    add_body(doc, 'O custo é calculado na segunda simulação (Etapa 3) com a curva de carga otimizada.')
    
    doc.add_heading('2.5 Pseudocódigo Completo', level=2)
    
    pseudo = """ALGORITMO: Análise 3 — Heurística com Deslizamento

ENTRADA: Microrrede (fontes, cargas, parâmetros)
SAÍDA: Despacho otimizado, custos, níveis de armazenamento

--- ETAPA 1: Simulação Original ---
1. curva_carga ← CurvaCarga(microrrede.carga)
2. custo_ordenado ← OrdenaFontesPorCusto(microrrede)
3. PARA t = 0 ATÉ 1439:
   3.1. D_rest ← curva_carga[t]
   3.2. excesso_sol ← 0
   3.3. PARA CADA fonte EM custo_ordenado:
        SE fonte = Solar:
           P_sol ← min(G_sol[t], D_rest)
           D_rest ← D_rest - P_sol
           excesso_sol ← G_sol[t] - P_sol
        SE fonte = Bateria:
           GerenciarBateria(D_rest, excesso_sol)
        SE fonte = Biogás:
           ProcessarBiogas(D_rest)
        SE fonte = Diesel:
           ProcessarDiesel(D_rest)
   3.4. ProcessarConcessionaria(D_rest)
   3.5. C[t] ← SomaCustosInstantaneos()

--- ETAPA 2: Deslizamento ---
4. curva_otimizada ← DeslizarCargas(curva_carga, C[t])

--- ETAPA 3: Simulação Otimizada ---
5. Repetir passos 1-3 com curva_otimizada
6. RETORNAR resultados"""
    
    p = doc.add_paragraph()
    run = p.add_run(pseudo)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    # ========================================
    # 3. ANÁLISE 5 — MILP
    # ========================================
    doc.add_page_break()
    doc.add_heading('3. Análise 5 — MILP com Deslizamento Integrado', level=1)
    
    doc.add_heading('3.1 Visão Geral', level=2)
    add_body(doc, 'A Análise 5 formula o problema de despacho econômico como um modelo de Programação Linear Inteira Mista (MILP — Mixed Integer Linear Programming). Diferentemente da Análise 3, todas as decisões — despacho de fontes, gerenciamento da bateria e deslizamento de cargas — são otimizadas simultaneamente em uma única resolução.')
    
    add_body(doc, 'Características principais:')
    add_bullet(doc, 'Otimização global: todas as decisões são tomadas conjuntamente')
    add_bullet(doc, 'Garantia de otimalidade: o solver encontra o mínimo global (ou prova infactibilidade)')
    add_bullet(doc, 'Deslizamento integrado: horários das cargas flexíveis são variáveis de decisão')
    add_bullet(doc, 'Restrições de exclusão mútua: bateria não carrega e descarrega simultaneamente')
    add_bullet(doc, 'Penalidade de suavização: evita oscilações rápidas na bateria')
    add_bullet(doc, 'Solver utilizado: CBC (COIN-OR Branch and Cut)')
    
    doc.add_heading('3.2 Variáveis de Decisão', level=2)
    
    table3 = doc.add_table(rows=1, cols=4, style='Table Grid')
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr3 = table3.rows[0].cells
    hdr3[0].text = 'Variável'
    hdr3[1].text = 'Tipo'
    hdr3[2].text = 'Domínio'
    hdr3[3].text = 'Descrição'
    for cell in hdr3:
        for paragraph in cell.paragraphs:
            paragraph.runs[0].bold = True
    
    vars_milp = [
        ('P_sol(t)', 'Contínua', '≥ 0', 'Potência solar utilizada no período t (kW)'),
        ('P_bat(t)', 'Contínua', '≥ 0', 'Potência de descarga da bateria no período t (kW)'),
        ('P_die(t)', 'Contínua', '≥ 0', 'Potência do diesel no período t (kW)'),
        ('P_conc(t)', 'Contínua', '≥ 0', 'Potência da concessionária no período t (kW)'),
        ('P_rec(t)', 'Contínua', '≥ 0', 'Potência de recarga da bateria no período t (kW)'),
        ('P_curtail(t)', 'Contínua', '≥ 0', 'Potência solar desperdiçada (curtailment) no período t (kW)'),
        ('E_bat(t)', 'Contínua', '≥ 0', 'Nível de energia da bateria no período t (kWh)'),
        ('E_die(t)', 'Contínua', '≥ 0', 'Nível do tanque diesel no período t (L)'),
        ('U_die(t)', 'Binária', '{0, 1}', '1 se diesel ligado no período t, 0 caso contrário'),
        ('U_bat(t)', 'Binária', '{0, 1}', '1 se bateria descarregando, 0 se carregando/inativa'),
        ('Δ_desc(t)', 'Contínua', '≥ 0', 'Variação absoluta da descarga da bateria entre t-1 e t'),
        ('Δ_rec(t)', 'Contínua', '≥ 0', 'Variação absoluta da recarga da bateria entre t-1 e t'),
        ('δₖ(s)', 'Binária', '{0, 1}', '1 se carga flexível k inicia no período s'),
    ]
    for var, tipo, dom, desc in vars_milp:
        row = table3.add_row().cells
        row[0].text = var
        row[1].text = tipo
        row[2].text = dom
        row[3].text = desc
    
    add_body(doc, f'Total de variáveis por período: 8 contínuas + 2 binárias = 10 variáveis × 1440 períodos + variáveis de deslizamento δₖ(s).')
    
    doc.add_heading('3.3 Função Objetivo', level=2)
    add_body(doc, 'O objetivo é minimizar o custo operacional total ao longo do dia:')
    
    add_equation(doc, 'min Z = C_die + C_bat + C_conc + C_sol + C_startup + C_suav + C_curtail')
    
    add_body(doc, 'Onde cada componente é definido como:')
    
    add_equation(doc, 'C_die = Σ_{t=0}^{1439} c_die × P_die(t) / 60')
    add_equation(doc, 'C_bat = Σ_{t=0}^{1439} c_bat × P_bat(t) / 60')
    add_equation(doc, 'C_conc = Σ_{t=0}^{1439} c_conc × P_conc(t) / 60')
    add_equation(doc, 'C_sol = 0.01 × Σ_{t=0}^{1439} P_sol(t)')
    add_equation(doc, 'C_startup = Σ_{t=1}^{1439} 50 × (U_die(t) − U_die(t−1))')
    add_equation(doc, 'C_suav = 0.5 × Σ_{t=1}^{1439} [Δ_desc(t) + Δ_rec(t)] / 60')
    add_equation(doc, 'C_curtail = 10 × Σ_{t=0}^{1439} P_curtail(t) / 60')
    
    add_body(doc, 'A divisão por 60 converte kW (potência) para kWh (energia) em intervalos de 1 minuto.')
    add_body(doc, 'A penalidade de curtailment (R$ 10/kWh) força o otimizador a usar toda a geração solar, armazenando o excesso na bateria em vez de desperdiçar.')
    add_body(doc, 'A penalidade de suavização (R$ 0.5/kWh) desencoraja mudanças rápidas na potência da bateria, produzindo curvas de operação mais estáveis e realistas.')
    
    doc.add_heading('3.4 Restrições', level=2)
    
    doc.add_heading('3.4.1 Balanço de Energia', level=3)
    add_body(doc, 'A oferta total deve ser igual à demanda total mais a recarga da bateria, para cada período t:')
    add_equation(doc, 'P_sol(t) + P_bat(t) + P_die(t) + P_conc(t) = D(t) + P_rec(t),   ∀t ∈ T')
    add_body(doc, 'Onde D(t) é a demanda total (fixa + flexível):')
    add_equation(doc, 'D(t) = D_fixo(t) + Σₖ pₖ × Σ_{s ∈ Ativos(k,t)} δₖ(s)')
    add_body(doc, 'Ativos(k,t) = {s ∈ Sₖ : s ≤ t < s + dₖ} é o conjunto de inícios que tornam a carga k ativa no período t.')
    
    doc.add_heading('3.4.2 Geração Solar', level=3)
    add_body(doc, 'Toda a geração solar deve ser contabilizada — utilizada ou curtailed:')
    add_equation(doc, 'P_sol(t) + P_curtail(t) = G_sol(t),   ∀t ∈ T')
    add_body(doc, 'Esta é uma restrição de igualdade: o otimizador não pode simplesmente ignorar a geração solar. Se não utilizar, deve pagar a penalidade de curtailment na função objetivo.')
    
    doc.add_heading('3.4.3 Limites do Diesel', level=3)
    add_body(doc, 'O diesel utiliza formulação Big-M com a binária U_die(t):')
    add_equation(doc, 'P_die(t) ≤ P_die_max × U_die(t),   ∀t ∈ T     [só opera se ligado]')
    add_equation(doc, 'P_die(t) ≥ 0.2 × P_die_max × U_die(t),   ∀t ∈ T     [carga mínima 20%]')
    add_body(doc, 'Quando U_die(t) = 0: P_die(t) = 0 (desligado)')
    add_body(doc, 'Quando U_die(t) = 1: 0.2 × P_die_max ≤ P_die(t) ≤ P_die_max (operando entre 20% e 100%)')
    
    doc.add_heading('3.4.4 Limites da Bateria e Exclusão Mútua', level=3)
    add_body(doc, 'Limites de potência:')
    add_equation(doc, 'P_bat(t) ≤ P_bat_max,   ∀t ∈ T     [limite de descarga]')
    add_equation(doc, 'P_rec(t) ≤ P_bat_max,   ∀t ∈ T     [limite de recarga]')
    
    add_body(doc, 'Exclusão mútua — a bateria não pode descarregar e carregar simultaneamente:')
    add_equation(doc, 'P_bat(t) ≤ P_bat_max × U_bat(t),   ∀t ∈ T')
    add_equation(doc, 'P_rec(t) ≤ P_bat_max × (1 − U_bat(t)),   ∀t ∈ T')
    add_body(doc, 'Quando U_bat(t) = 1: bateria pode descarregar (P_bat ≤ P_max, P_rec = 0)')
    add_body(doc, 'Quando U_bat(t) = 0: bateria pode carregar (P_bat = 0, P_rec ≤ P_max)')
    
    doc.add_heading('3.4.5 Dinâmica de Armazenamento da Bateria', level=3)
    add_body(doc, 'A evolução do estado de carga segue:')
    add_equation(doc, 'E_bat(0) = E_bat_0     [estado inicial = capacidade máxima]')
    add_equation(doc, 'E_bat(t+1) = E_bat(t) − P_bat(t)/60 + η × P_rec(t)/60,   ∀t ∈ T')
    add_equation(doc, 'E_bat_min ≤ E_bat(t) ≤ E_bat_max,   ∀t ∈ T')
    add_body(doc, 'A descarga reduz o nível diretamente (sem fator de eficiência na saída), enquanto a recarga multiplica pela eficiência η (perda na conversão).')
    
    doc.add_heading('3.4.6 Suavização da Bateria', level=3)
    add_body(doc, 'Para evitar oscilações rápidas, variáveis auxiliares capturam o valor absoluto das mudanças:')
    add_equation(doc, 'Δ_desc(t) ≥ P_bat(t) − P_bat(t−1),   ∀t ≥ 1')
    add_equation(doc, 'Δ_desc(t) ≥ P_bat(t−1) − P_bat(t),   ∀t ≥ 1')
    add_equation(doc, 'Δ_rec(t) ≥ P_rec(t) − P_rec(t−1),   ∀t ≥ 1')
    add_equation(doc, 'Δ_rec(t) ≥ P_rec(t−1) − P_rec(t),   ∀t ≥ 1')
    add_body(doc, 'Estas restrições garantem que Δ_desc(t) ≥ |P_bat(t) − P_bat(t−1)| e Δ_rec(t) ≥ |P_rec(t) − P_rec(t−1)|. Como a função objetivo minimiza esses valores, na solução ótima a igualdade é atingida.')
    
    doc.add_heading('3.4.7 Dinâmica de Armazenamento do Diesel', level=3)
    add_equation(doc, 'E_die(0) = E_die_0     [tanque inicial cheio]')
    add_equation(doc, 'E_die(t+1) = E_die(t) − σ × P_die(t)/60,   ∀t ∈ T')
    add_equation(doc, '0 ≤ E_die(t) ≤ E_die_max,   ∀t ∈ T')
    
    doc.add_heading('3.4.8 Deslizamento Integrado de Cargas', level=3)
    add_body(doc, 'Para cada carga flexível k ∈ K:')
    add_equation(doc, 'Σ_{s ∈ Sₖ} δₖ(s) = 1     [exatamente um horário de início]')
    add_body(doc, 'O conjunto Sₖ contém os possíveis inícios com passo de 15 minutos:')
    add_equation(doc, 'Sₖ = {0, 15, 30, ..., 1440 − dₖ}')
    add_body(doc, 'A demanda variável é automaticamente incorporada no balanço de energia:')
    add_equation(doc, 'D(t) = D_fixo(t) + Σₖ pₖ × Σ_{s ∈ Ativos(k,t)} δₖ(s)')
    add_body(doc, 'O otimizador determina simultaneamente o melhor horário de cada carga flexível e o despacho ótimo de todas as fontes.')
    
    # ========================================
    # 4. COMPARAÇÃO
    # ========================================
    doc.add_page_break()
    doc.add_heading('4. Comparação entre Análise 3 e Análise 5', level=1)
    
    doc.add_heading('4.1 Tabela Comparativa', level=2)
    
    table4 = doc.add_table(rows=1, cols=3, style='Table Grid')
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr4 = table4.rows[0].cells
    hdr4[0].text = 'Aspecto'
    hdr4[1].text = 'Análise 3 (Heurística)'
    hdr4[2].text = 'Análise 5 (MILP)'
    for cell in hdr4:
        for paragraph in cell.paragraphs:
            paragraph.runs[0].bold = True
    
    comparacao = [
        ('Método', 'Heurística de cascata + deslizamento em 2 fases', 'Programação Linear Inteira Mista (MILP)'),
        ('Solver', 'Algoritmo sequencial em Python', 'CBC (COIN-OR Branch and Cut)'),
        ('Otimalidade', 'Solução viável, não necessariamente ótima', 'Solução ótima global garantida'),
        ('Deslizamento', '2 fases: simula → desliza → re-simula', 'Integrado: uma única otimização'),
        ('Variáveis binárias', 'Nenhuma', 'U_die(t), U_bat(t), δₖ(s)'),
        ('Ordenação de fontes', 'Custo crescente (cascata sequencial)', 'Decidido pelo otimizador (custo na FO)'),
        ('Bateria', 'Descarga primeiro, recarga do excesso solar depois', 'Descarga/recarga otimizada globalmente'),
        ('Exclusão carga/descarga', 'Implícita (lógica if/else)', 'Explícita (U_bat binária)'),
        ('Suavização', 'Não possui', 'Penalidade por mudanças rápidas'),
        ('Curtailment solar', 'Excesso não armazenado é perdido', 'Penalidade força aproveitamento'),
        ('Custo de startup', 'Não modelado', 'R$ 50 por partida do diesel'),
        ('Complexidade', 'O(n) — linear no número de períodos', 'NP-hard (Branch & Cut)'),
        ('Tempo de execução', 'Milissegundos', 'Segundos a minutos'),
        ('Escalabilidade', 'Excelente', 'Limitada por variáveis binárias'),
    ]
    for asp, a3, a5 in comparacao:
        row = table4.add_row().cells
        row[0].text = asp
        row[1].text = a3
        row[2].text = a5
    
    doc.add_heading('4.2 Vantagens e Desvantagens', level=2)
    
    doc.add_heading('Análise 3 — Heurística', level=3)
    add_body(doc, 'Vantagens:')
    add_bullet(doc, 'Execução extremamente rápida (milissegundos)')
    add_bullet(doc, 'Intuitiva e fácil de interpretar/depurar')
    add_bullet(doc, 'Escalável para muitas fontes e cargas')
    add_bullet(doc, 'Não requer software de otimização especializado')
    
    add_body(doc, 'Desvantagens:')
    add_bullet(doc, 'Não garante a solução ótima global')
    add_bullet(doc, 'Deslizamento em duas fases pode ser subótimo (primeiro fixa custos, depois move cargas)')
    add_bullet(doc, 'Não considera interdependências temporais complexas')
    add_bullet(doc, 'Sem modelagem de custos de startup/shutdown')
    
    doc.add_heading('Análise 5 — MILP', level=3)
    add_body(doc, 'Vantagens:')
    add_bullet(doc, 'Garante solução ótima global (mínimo custo)')
    add_bullet(doc, 'Otimização simultânea de despacho e deslizamento')
    add_bullet(doc, 'Modelo fisicamente mais realista (exclusão mútua, suavização, startup)')
    add_bullet(doc, 'Flexível para adicionar novas restrições')
    
    add_body(doc, 'Desvantagens:')
    add_bullet(doc, 'Tempo de resolução maior (segundos a minutos)')
    add_bullet(doc, 'Complexidade computacional NP-hard')
    add_bullet(doc, 'Requer solver MILP (PuLP/CBC)')
    add_bullet(doc, 'Número de variáveis binárias cresce com cargas flexíveis')
    
    doc.add_heading('4.3 Resumo dos Modelos Matemáticos', level=2)
    
    add_body(doc, 'Análise 3 — Modelo Heurístico:')
    add_equation(doc, 'min C = Σ_t [Σ_f c_f × P_f(t) / 60]')
    add_equation(doc, 'sujeito a: despacho sequencial por custo crescente')
    add_equation(doc, '           deslizamento em 2 fases (ex-post)')
    
    add_body(doc, 'Análise 5 — Modelo MILP:')
    add_equation(doc, 'min Z = Σ_t [c_die×P_die(t) + c_bat×P_bat(t) + c_conc×P_conc(t)]/60')
    add_equation(doc, '       + 0.01×Σ P_sol + 50×Σ startup_die')
    add_equation(doc, '       + 0.5×Σ [Δ_desc(t) + Δ_rec(t)]/60')
    add_equation(doc, '       + 10×Σ P_curtail(t)/60')
    add_equation(doc, 'sujeito a: balanço de energia (igualdade)')
    add_equation(doc, '           limites de potência e armazenamento')
    add_equation(doc, '           exclusão mútua carga/descarga bateria')
    add_equation(doc, '           dinâmica de armazenamento (bateria + diesel)')
    add_equation(doc, '           suavização da bateria')
    add_equation(doc, '           deslizamento integrado (δₖ(s) binárias)')
    
    # ========================================
    # SALVAR
    # ========================================
    output_path = os.path.join(os.path.dirname(__file__), 'Formulacao_Matematica_Analises_3_e_5.docx')
    doc.save(output_path)
    print(f'Documento salvo em: {output_path}')
    return output_path

if __name__ == '__main__':
    gerar_documento()
